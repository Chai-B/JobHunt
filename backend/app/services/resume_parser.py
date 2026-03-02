import io
import re
import unicodedata
import fitz  # PyMuPDF
import docx
from sentence_transformers import SentenceTransformer
from loguru import logger

# Lazy-loaded at module level for worker reuse
model = SentenceTransformer('all-MiniLM-L6-v2')

# Resume section headers — used for both scoring and text segmentation
_SECTION_HEADERS = frozenset({
    "experience", "education", "skills", "projects", "summary", "objective",
    "certifications", "awards", "publications", "interests", "references",
    "work history", "professional experience", "technical skills",
    "core competencies", "achievements", "volunteer",
})

# Action verbs for semantic quality scoring
_ACTION_VERBS = frozenset({
    "developed", "managed", "led", "built", "designed", "implemented",
    "created", "improved", "achieved", "delivered", "launched", "optimized",
    "architected", "deployed", "integrated", "automated", "scaled",
    "mentored", "coordinated", "analyzed", "resolved", "streamlined",
    "collaborated", "engineered", "spearheaded", "pioneered",
})


def _normalize_text(text: str) -> str:
    """Clean extracted text: normalize unicode, collapse whitespace, strip noise."""
    # Normalize unicode (e.g. ligatures, special chars)
    text = unicodedata.normalize("NFKD", text)
    
    # Replace common PDF artifacts
    text = text.replace('\x00', '')  # Null bytes
    text = text.replace('\uf0b7', '•')  # Common bullet replacement char
    text = text.replace('\uf0a7', '•')
    
    # Collapse excessive whitespace within lines (preserve newlines)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = re.sub(r'[ \t]+', ' ', line).strip()
        if line:
            cleaned_lines.append(line)
    
    # Collapse excessive blank lines (max 2 consecutive)
    result = []
    blank_count = 0
    for line in cleaned_lines:
        if not line:
            blank_count += 1
            if blank_count <= 2:
                result.append(line)
        else:
            blank_count = 0
            result.append(line)
    
    return '\n'.join(result).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract and clean text from a PDF byte stream."""
    text_parts = []
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                    
        if not text_parts:
            # Fallback: try extracting with different method (handles some scanned PDFs better)
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    page_text = page.get_text("blocks")
                    if page_text:
                        block_text = "\n".join(
                            block[4] for block in page_text 
                            if block[6] == 0 and isinstance(block[4], str) and block[4].strip()
                        )
                        if block_text.strip():
                            text_parts.append(block_text)
                            
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")
    
    raw = "\n".join(text_parts)
    return _normalize_text(raw)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract and clean text from a DOCX byte stream."""
    text_parts = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
                    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    raw = "\n".join(text_parts)
    return _normalize_text(raw)


def _compute_structural_score(text: str) -> float:
    """Score how well-structured this resume is (0.0 - 1.0)."""
    score = 0.0
    text_lower = text.lower()
    
    # Section coverage (up to 0.4)
    found_sections = sum(1 for s in _SECTION_HEADERS if s in text_lower)
    score += min(found_sections / 4.0, 1.0) * 0.4
    
    # Bullet points / structured lists (up to 0.3)
    bullet_count = len(re.findall(r'[•\-\*▪►]\s', text))
    score += min(bullet_count / 10.0, 1.0) * 0.3
    
    # Reasonable word count (up to 0.2)
    word_count = len(text.split())
    if 200 <= word_count <= 2000:
        score += 0.2
    elif word_count > 80:
        score += 0.1
    
    # Consistent formatting (up to 0.1) — check for dates which indicate structured content
    date_count = len(re.findall(r'(20\d{2}|19\d{2})\s*[-–—]\s*(20\d{2}|19\d{2}|present|current)', text_lower))
    score += min(date_count / 3.0, 1.0) * 0.1
    
    return round(min(score, 1.0), 2)


def _compute_semantic_score(text: str) -> float:
    """Score how content-rich this resume is (0.0 - 1.0)."""
    text_lower = text.lower()
    
    # Action verbs (up to 0.4)
    verb_count = sum(1 for v in _ACTION_VERBS if v in text_lower)
    verb_score = min(verb_count / 6.0, 1.0) * 0.4
    
    # Quantified achievements (up to 0.3)
    quant_patterns = [
        r'\d+[%+]',                           # Percentages / growth
        r'\$[\d,]+[KMB]?',                    # Dollar amounts
        r'\d+\s*(users|customers|clients)',    # User counts
        r'\d+\s*(projects|applications)',      # Project counts
        r'\d+x\s',                            # Multipliers
    ]
    quant_count = sum(len(re.findall(p, text_lower)) for p in quant_patterns)
    quant_score = min(quant_count / 4.0, 1.0) * 0.3
    
    # Technical depth — mentions of specific technologies (up to 0.3)
    tech_keywords = [
        "python", "javascript", "react", "aws", "docker", "kubernetes",
        "sql", "typescript", "node", "java", "api", "microservices",
        "ci/cd", "agile", "machine learning", "deep learning",
    ]
    tech_count = sum(1 for t in tech_keywords if t in text_lower)
    tech_score = min(tech_count / 5.0, 1.0) * 0.3
    
    return round(min(verb_score + quant_score + tech_score, 1.0), 2)


def parse_and_embed_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Full resume processing pipeline:
    1. Extract text based on file extension
    2. Normalize and clean the text
    3. Generate sentence embedding for semantic matching
    4. Score structural and semantic quality
    
    Returns dict with raw_text, embedding, structural_score, semantic_score.
    Raises ValueError if extraction produces no usable text.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw_text = extract_text_from_docx(file_bytes)
    elif ext in ("txt", "md"):
        raw_text = file_bytes.decode('utf-8', errors='ignore')
        raw_text = _normalize_text(raw_text)
    else:
        raise ValueError(f"Unsupported file format: .{ext}")
        
    if not raw_text or len(raw_text.strip()) < 50:
        raise ValueError(
            f"Extraction produced insufficient text ({len(raw_text.strip())} chars). "
            f"The file may be image-only, corrupted, or password-protected."
        )

    logger.info(f"Extracted {len(raw_text)} chars from {filename}")
    
    # Truncate for embedding — MiniLM-L6-v2 has ~256 token window, 8000 chars is generous
    embed_text = raw_text[:8000]
    embedding = model.encode(embed_text).tolist()
    
    structural_score = _compute_structural_score(raw_text)
    semantic_score = _compute_semantic_score(raw_text)
    
    logger.info(f"Resume {filename}: structural={structural_score}, semantic={semantic_score}")
    
    return {
        "raw_text": raw_text,
        "embedding": embedding,
        "structural_score": structural_score,
        "semantic_score": semantic_score,
    }
